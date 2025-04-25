import json
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class Exporter:
    @staticmethod
    def save_json(data, filename="output.json"):
        try:
            with open(filename, 'w') as f:
                json.dump(data, f, indent=4)
            print(f"[✅] JSON saved to {filename}")
        except Exception as e:
            print(f"[❌] Error saving JSON: {e}")

    @staticmethod
    def save_csv(data, filename="output.csv"):
        try:
            with open(filename, 'w', newline='') as f:
                writer = csv.writer(f)
                writer.writerow(data[0].keys())  # Writing header
                for row in data:
                    writer.writerow(row.values())
            print(f"[✅] CSV saved to {filename}")
        except Exception as e:
            print(f"[❌] Error saving CSV: {e}")

    @staticmethod
    def to_word(content_text, image_path=None, sources=None, filename="report.docx"):
        try:
            doc = Document()
            
            # Title and metadata
            title = doc.add_heading("AI Research Report", level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            date_para = doc.add_paragraph()
            date_para.add_run("Generated on: ").bold = True
            date_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            # Add content with improved formatting
            doc.add_heading("Research Findings", level=1)
            for paragraph in content_text.split('\n'):
                if paragraph.strip().startswith('#'):
                    # Handle markdown headings
                    level = paragraph.count('#')
                    doc.add_heading(paragraph.replace('#', '').strip(), level=min(level, 3))
                elif paragraph.strip():
                    p = doc.add_paragraph(paragraph)
                    p.paragraph_format.space_after = Pt(8)
            
            # Add visualization if available
            if image_path:
                doc.add_heading("Source Reliability Analysis", level=1)
                doc.add_picture(image_path, width=Inches(6))
            
            # Add sources section if available
            if sources:
                doc.add_heading("References", level=1)
                for i, source in enumerate(sources, 1):
                    if isinstance(source, dict):
                        url = source.get('url', '')
                        title = source.get('title', f'Source {i}')
                    else:
                        url = source
                        title = f'Source {i}'
                    
                    p = doc.add_paragraph()
                    p.add_run(f"[{i}] ").bold = True
                    p.add_run(f"{title}: {url}")
            
            doc.save(filename)
            print(f"[✅] Report saved to {filename}")
            return filename
        except Exception as e:
            print(f"[❌] Error saving report: {e}")
            return None

    @staticmethod
    def format_output(data, format_type="json"):
        if format_type == "json":
            Exporter.save_json(data)
        elif format_type == "csv":
            Exporter.save_csv(data)
        else:
            print(f"[❌] Unsupported format: {format_type}")
