import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import datetime
from docx import Document
from docx.shared import Inches
from typing import List, Dict, Optional
import os
import logging

class VisualizationAgent:
    def __init__(self):
        self.chart_files = []
        self.logger = logging.getLogger(__name__)
        try:
            # Updated style setting that works with modern matplotlib
            plt.style.use('seaborn-v0_8')
            self.logger.info("Visualization styles initialized")
        except Exception as e:
            self.logger.warning(f"Could not set seaborn style: {str(e)}")
            plt.style.use('default')

    def plot_reliability(self, sources: List[Dict], title: str = "Source Reliability", 
                         xlabel: str = "Sources", ylabel: str = "Reliability Score", 
                         filename: str = "reliability_plot.png") -> Optional[str]:
        """Generate a reliability score bar chart"""
        try:
            if not sources:
                raise ValueError("No sources provided")
                
            # Prepare data
            labels = []
            scores = []
            colors = []
            
            for i, src in enumerate(sources):
                title = src.get("title", f"Source {i+1}")
                labels.append(title[:25] + "..." if len(title) > 25 else title)
                score = float(src.get("score", 0.0))
                scores.append(score)
                colors.append(self._get_score_color(score))
            
            # Create figure with improved layout
            plt.figure(figsize=(12, 6))
            bars = plt.barh(labels, scores, color=colors)
            plt.title(title, pad=20)
            plt.xlabel(ylabel)
            plt.xlim(0, 1.0)
            
            # Add score labels
            for bar in bars:
                width = bar.get_width()
                plt.text(width + 0.02, bar.get_y() + bar.get_height()/2,
                         f'{width:.2f}', ha='left', va='center')
            
            plt.tight_layout()
            plt.savefig(filename, dpi=300, bbox_inches='tight')
            plt.close()
            
            self.chart_files.append(filename)
            self.logger.info(f"Generated reliability plot: {filename}")
            return filename
        except Exception as e:
            self.logger.error(f"Error generating reliability plot: {str(e)}")
            return None

    def _get_score_color(self, score: float) -> str:
        """Get color based on reliability score"""
        if score > 0.8:
            return '#2ecc71'  # Green
        elif score > 0.6:
            return '#f39c12'  # Orange
        else:
            return '#e74c3c'  # Red

    def save_visualizations_to_docx(self, filename: str = "visualizations.docx") -> bool:
        """Save all generated charts to a Word document"""
        try:
            if not self.chart_files:
                self.logger.warning("No visualization files to save")
                return False
                
            doc = Document()
            doc.add_heading("Research Visualizations", level=1)
            doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

            for chart_file in self.chart_files:
                try:
                    doc.add_heading(os.path.basename(chart_file), level=2)
                    doc.add_picture(chart_file, width=Inches(6))
                    doc.add_page_break()
                except Exception as e:
                    self.logger.error(f"Error adding {chart_file}: {str(e)}")
                    continue
            
            doc.save(filename)
            self.logger.info(f"Visualizations saved to {filename}")
            return True
        except Exception as e:
            self.logger.error(f"Error saving visualizations: {str(e)}")
            return False
